# Version 5

# Export_Agent.py (Final Fixed Version)

import re
import io
import time
import json
import threading
import requests
from datetime import datetime

from docx import Document
from docx.shared import Pt as DocxPt, RGBColor as DocxRGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from azure.storage.blob import BlobServiceClient

# ──────────────────────────────────────────────────────────────
# Config
# ──────────────────────────────────────────────────────────────
BLOB_CONFIG = {
    "account_url": "https://cxqaazureaihub8779474245.blob.core.windows.net",
    "sas_token": "sv=2022-11-02&ss=bfqt&srt=sco&sp=rwdlacupiytfx&se=2030-11-21T02:02:26Z&st=2024-11-20T18:02:26Z&spr=https&sig=YfZEUMeqiuBiG7le2JfaaZf%2FW6t8ZW75yCsFM6nUmUw%3D",
    "container": "5d74a98c-1fc6-4567-8545-2632b489bd0b-azureml-blobstore"
}

# ──────────────────────────────────────────────────────────────
# Retry-safe OpenAI call
# ──────────────────────────────────────────────────────────────
def openai_call_with_retry(endpoint, headers, payload, max_attempts=3, backoff=5, timeout=30):
    attempts = 0
    while attempts < max_attempts:
        try:
            response = requests.post(endpoint, headers=headers, json=payload, timeout=timeout)
            response.raise_for_status()
            return response.json()
        except Exception as e:
            attempts += 1
            if attempts >= max_attempts:
                return {"error": f"API_ERROR: {str(e)}"}
            time.sleep(backoff)

# ──────────────────────────────────────────────────────────────
# Upload to Azure Blob and get download link
# ──────────────────────────────────────────────────────────────
def upload_to_azure_blob(blob_config, file_buffer, file_name_prefix):
    try:
        blob_service = BlobServiceClient(
            account_url=blob_config["account_url"],
            credential=blob_config["sas_token"]
        )
        container_client = blob_service.get_container_client(blob_config["container"])
        file_name = f"{file_name_prefix}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
        blob_client = container_client.get_blob_client(file_name)
        blob_client.upload_blob(file_buffer, overwrite=True)
        url = f"{blob_config['account_url']}/{blob_config['container']}/{file_name}?{blob_config['sas_token']}"
        threading.Timer(300, blob_client.delete_blob).start()
        return url
    except Exception as e:
        raise Exception(f"Azure Blob Upload Error: {str(e)}")

# ──────────────────────────────────────────────────────────────
# Generate Document
# ──────────────────────────────────────────────────────────────
def Call_DOC(latest_question, latest_answer, chat_history, instructions):
    def generate_doc_text():
        prompt = f"""You are a professional document writer. Use ONLY the provided information to generate a document.
Format:
Section Heading
- Bullet 1
- Bullet 2

Separate sections with double newlines.

Data:
- Instructions: {instructions}
- Question: {latest_question}
- Answer: {latest_answer}
- Chat History: {chat_history}
"""
        endpoint = "https://cxqaazureaihub2358016269.openai.azure.com/openai/deployments/gpt-4o/chat/completions?api-version=2025-01-01-preview"
        headers = {
            "Content-Type": "application/json",
            "api-key": "Cv54PDKaIusK0dXkMvkBbSCgH982p1CjUwaTeKlir1NmB6tycSKMJQQJ99AKACYeBjFXJ3w3AAAAACOGllor"
        }
        payload = {
            "messages": [
                {"role": "system", "content": "Generate structured document content"},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 1000,
            "temperature": 0.3
        }
        result = openai_call_with_retry(endpoint, headers, payload)
        if "error" in result:
            return f"API_ERROR: {result['error']}"
        try:
            return result["choices"][0]["message"]["content"].strip()
        except Exception as e:
            return f"API_ERROR: {str(e)}"

    doc_text = generate_doc_text()
    if doc_text.startswith("API_ERROR:"):
        return f"OpenAI API Error: {doc_text[10:]}"
    if "NOT_ENOUGH_INFO" in doc_text.upper():
        return "Error: Insufficient information to generate document"
    if len(doc_text) < 20:
        return "Error: Generated content too short or invalid"

    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Cairo"
        style.font.size = DocxPt(12)
        style.font.color.rgb = DocxRGBColor(0, 0, 0)

        for section in doc.sections:
            sectPr = section._sectPr
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EAD7C2"/>')
            sectPr.append(shd)

        for section in doc_text.split("\n\n"):
            lines = [line.strip() for line in section.split("\n") if line.strip()]
            if not lines:
                continue
            heading = doc.add_heading(level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = heading.add_run(lines[0])
            run.font.color.rgb = DocxRGBColor(193, 114, 80)
            run.font.size = DocxPt(16)
            run.bold = True
            for bullet in lines[1:]:
                para = doc.add_paragraph(style="ListBullet")
                para.add_run(bullet.replace("- ", "").strip())

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        url = upload_to_azure_blob(BLOB_CONFIG, buffer, "document")
        return f"Here is your generated Document:\n{url}".strip()
    except Exception as e:
        return f"Document Generation Error: {str(e)}"

# ──────────────────────────────────────────────────────────────
# Call_Export dispatcher (add other formats as needed)
# ──────────────────────────────────────────────────────────────
def Call_Export(latest_question, latest_answer, chat_history, instructions):
    instructions_lower = instructions.lower()
    if re.search(r"\b(document|report|docx|white\s?paper|summary|contract|memo|manual)\b", instructions_lower):
        return Call_DOC(latest_question, latest_answer, chat_history, instructions)
    return "Not enough Information to perform export."
